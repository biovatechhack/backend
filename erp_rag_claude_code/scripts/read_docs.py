#!/usr/bin/env python3
"""
read_docs.py — Extract and cache content from source-of-truth DOCX files.
Run at the start of every Claude Code session.

Usage:
    python3 scripts/read_docs.py
    python3 scripts/read_docs.py --section "SQL Pipeline"
    python3 scripts/read_docs.py --sprint 4
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

DOCS_DIR = Path(__file__).parent.parent / "docs"
OUTPUT_FILE = DOCS_DIR / "source_of_truth.md"

MAPPING_DOC  = DOCS_DIR / "ERP_RAG_Architecture_Mapping.docx"
SPRINT_DOC   = DOCS_DIR / "ERP_RAG_Sprint_Plan_GitStrategy.docx"


def extract_docx(path: Path) -> str:
    """Extract all text from a .docx file."""
    try:
        from docx import Document  # type: ignore
    except ImportError:
        print("Installing python-docx...", file=sys.stderr)
        import subprocess
        subprocess.run(
            [sys.executable, "-m", "pip", "install", "python-docx", "-q"],
            check=True
        )
        from docx import Document

    doc = Document(str(path))
    lines = []

    for element in doc.element.body:
        # Paragraphs
        if element.tag.endswith('}p'):
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para and para.text.strip():
                style = para.style.name if para.style else ""
                if style.startswith("Heading 1"):
                    lines.append(f"\n## {para.text.strip()}")
                elif style.startswith("Heading 2"):
                    lines.append(f"\n### {para.text.strip()}")
                elif style.startswith("Heading 3"):
                    lines.append(f"\n#### {para.text.strip()}")
                else:
                    lines.append(para.text.strip())

        # Tables
        elif element.tag.endswith('}tbl'):
            for table in doc.tables:
                if table._element is element:
                    for row in table.rows:
                        cells = [c.text.strip() for c in row.cells]
                        # Deduplicate merged cells
                        seen = []
                        for c in cells:
                            if not seen or c != seen[-1]:
                                seen.append(c)
                        line = " | ".join(seen)
                        if line.strip(" |"):
                            lines.append(line)
                    lines.append("")
                    break

    return "\n".join(lines)


def build_source_of_truth() -> str:
    """Build the combined source-of-truth markdown."""
    sections = []

    sections.append("# ERP Agentic RAG — Source of Truth")
    sections.append(f"_Auto-generated from DOCX files. Do not edit manually._\n")

    # ── Architecture Mapping ──────────────────────────────────────
    sections.append("---\n# DOCUMENT 1: Architecture Mapping\n---\n")
    if MAPPING_DOC.exists():
        print(f"Reading: {MAPPING_DOC.name}...", file=sys.stderr)
        mapping_text = extract_docx(MAPPING_DOC)
        sections.append(mapping_text)
        # Add critical facts summary
        sections.append("""
---
## CRITICAL FACTS (Architecture Mapping)

1. **MongoDB** is the primary RAG store — NOT PostgreSQL
   - PostgreSQL is ERP read-only target only (via query_executor.py)
2. **SQL pipeline is 3 stages** — never collapse to 2:
   - Stage 1: query_generator.py → SQLGenerationResult
   - Stage 2: query_validator.py → ValidationReport (NO LLM call, pure code)
   - Stage 3: query_executor.py → ExecutionResult (only if is_valid=True AND has_tenant_filter=True)
3. **ModuleAccessGuard** is embedded in RBACMiddleware — NOT a separate file
4. **Middleware order**: Logging → Auth → RateLimit → RBAC → PIIMasking
5. **Two missing modules** to create: observability/ and evaluation/
6. **Celery workers** handle async ingestion — not in spec but in code
7. **vLLM** is a fallback to OpenAI — both clients exist
8. **MinIO** is production file storage — local is dev only
---
""")
    else:
        sections.append(f"⚠️  FILE NOT FOUND: {MAPPING_DOC}")
        sections.append(f"Place the file at: {MAPPING_DOC}\n")

    # ── Sprint Plan ───────────────────────────────────────────────
    sections.append("---\n# DOCUMENT 2: Sprint Plan & Git Strategy\n---\n")
    if SPRINT_DOC.exists():
        print(f"Reading: {SPRINT_DOC.name}...", file=sys.stderr)
        sprint_text = extract_docx(SPRINT_DOC)
        sections.append(sprint_text)
        # Add git quick-ref
        sections.append("""
---
## GIT QUICK REFERENCE

### Branch names
- Sprint:   sprint-N/theme          (e.g. sprint-1/observability)
- Feature:  feature/short-desc      (e.g. feature/prometheus-metrics)
- Hotfix:   hotfix/short-desc
- Experiment: experiment/desc       (NEVER merges to develop)

### Commit types
feat | fix | test | refactor | docs | chore | perf | security

### Sprint close sequence
1. git merge --no-ff sprint-N/theme → develop
2. git tag -a sprint-N-done -m "Sprint N: [summary]"
3. git push origin develop sprint-N-done
4. git push origin --delete sprint-N/theme

### Release sequence (Sprint 10)
1. git merge --no-ff develop → main
2. git tag -a v1.0.0 -m "Release v1.0.0"
3. git push origin main --tags
---
""")
    else:
        sections.append(f"⚠️  FILE NOT FOUND: {SPRINT_DOC}")
        sections.append(f"Place the file at: {SPRINT_DOC}\n")

    return "\n".join(sections)


def filter_section(content: str, query: str) -> str:
    """Return only sections matching the query string."""
    lines = content.split("\n")
    result = []
    in_section = False
    for line in lines:
        if query.lower() in line.lower():
            in_section = True
        elif line.startswith("## ") and in_section:
            in_section = False
        if in_section:
            result.append(line)
    return "\n".join(result) if result else f"No section matching '{query}' found."


def filter_sprint(content: str, sprint_num: int) -> str:
    """Return only the sprint N section."""
    patterns = [
        f"SPRINT {sprint_num} ",
        f"Sprint {sprint_num} —",
        f"Sprint {sprint_num}:",
    ]
    lines = content.split("\n")
    result = []
    in_sprint = False
    for line in lines:
        if any(p in line for p in patterns):
            in_sprint = True
        elif line.startswith("## ") and in_sprint and not any(p in line for p in patterns):
            in_sprint = False
        if in_sprint:
            result.append(line)
    return "\n".join(result) if result else f"Sprint {sprint_num} section not found."


def main():
    parser = argparse.ArgumentParser(description="Read source-of-truth DOCX files")
    parser.add_argument("--section", help="Filter to section containing this text")
    parser.add_argument("--sprint", type=int, help="Filter to specific sprint number")
    parser.add_argument("--no-cache", action="store_true", help="Force re-extraction")
    args = parser.parse_args()

    # Build or load cache
    if OUTPUT_FILE.exists() and not args.no_cache:
        print(f"✓ Loading cached source of truth: {OUTPUT_FILE}", file=sys.stderr)
        content = OUTPUT_FILE.read_text(encoding="utf-8")
    else:
        print("Extracting from DOCX files...", file=sys.stderr)
        content = build_source_of_truth()
        OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
        OUTPUT_FILE.write_text(content, encoding="utf-8")
        print(f"✓ Saved to: {OUTPUT_FILE}", file=sys.stderr)

    # Apply filters
    if args.section:
        content = filter_section(content, args.section)
    elif args.sprint:
        content = filter_sprint(content, args.sprint)

    print(content)
    print(f"\n✓ Source of truth loaded. ({len(content)} chars)", file=sys.stderr)


if __name__ == "__main__":
    main()
